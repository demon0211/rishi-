import sys
import docx

doc = docx.Document(r"c:\Users\gokul\OneDrive\Desktop\rishi 2\backend\outputs\test_output.docx")
tables = doc.tables
if not tables:
    print("No tables found")
    sys.exit()

table = tables[0]
for r_idx, row in enumerate(table.rows):
    print(f"Row {r_idx}")
    for c_idx, cell in enumerate(row.cells):
        print(f"  Cell {c_idx}")
        for p_idx, p in enumerate(cell.paragraphs):
            runs = p.runs
            text = p.text
            if not text.strip(): continue
            format_info = []
            for run in runs:
                if run.text.strip():
                     format_info.append(f"[{run.text}] Italic: {run.italic}, Bold: {run.bold}, Size: {run.font.size.pt if run.font.size else None}")
            print(f"    Para {p_idx}: {text}")
            for info in format_info:
                print(f"      {info}")
