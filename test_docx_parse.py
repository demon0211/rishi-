import sys
sys.path.append(r"c:\Users\gokul\OneDrive\Desktop\rishi 2\backend")
from nlp_processor import NLPProcessor
from formatter import WordGenerator
import docx

nlp_processor = NLPProcessor()
raw_text = nlp_processor.extract_text_from_file(r"c:\Users\gokul\OneDrive\Desktop\rishi 2\backend\uploads\CKD_ML_Prediction_Enhanced_Paper_111.docx")
doc_data = nlp_processor.process_text(raw_text)

word_gen = WordGenerator(r"c:\Users\gokul\OneDrive\Desktop\rishi 2\backend\outputs")
word_gen.generate_docx(doc_data, "debug_ckd_output.docx")

doc = docx.Document(r"c:\Users\gokul\OneDrive\Desktop\rishi 2\backend\outputs\debug_ckd_output.docx")
tables = doc.tables
if not tables:
    print("No tables generated for authors")
else:
    table = tables[0]
    for r_idx, row in enumerate(table.rows):
        print(f"Row {r_idx}")
        for c_idx, cell in enumerate(row.cells):
            print(f"  Cell {c_idx}")
            for p_idx, p in enumerate(cell.paragraphs):
                text = p.text
                if text.strip():
                    print(f"    Para {p_idx}: {text}")

