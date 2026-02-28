"""
IEEE-Standard Document Formatter — Complete Redesign
Produces PDF (via ReportLab BaseDocTemplate + Frames) and DOCX (via python-docx)
outputs that match an IEEE conference paper layout.

DocumentBuilder pattern:
  build_title() → build_authors() → build_abstract() → build_keywords()
  → switch_to_two_column() → build_sections() → build_references()
"""

import os
import textwrap
import html
import re
from typing import List

# ── ReportLab ─────────────────────────────────────────────────────────────
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import (
    BaseDocTemplate, PageTemplate, Frame,
    Paragraph, Spacer, FrameBreak, NextPageTemplate, HRFlowable,
    KeepTogether
)
from reportlab.platypus.tableofcontents import SimpleIndex
from reportlab.lib import colors

# ── python-docx ───────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT # Added for centering tables
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Internal ──────────────────────────────────────────────────────────────
from nlp_processor import DocumentData, AuthorInfo, SectionData


# ═══════════════════════════════════════════════════════════════════════════
# SHARED CONSTANTS (IEEE A4 Specs)
# ═══════════════════════════════════════════════════════════════════════════

CM = 28.346 # ReportLab points per cm
COL_W = 8.25 * CM
COL_GAP = 0.8 * CM
LEFT_MARGIN = 1.9 * CM
RIGHT_MARGIN = 1.9 * CM
TOP_MARGIN = 2.54 * CM
BOTTOM_MARGIN = 2.86 * CM


# ═══════════════════════════════════════════════════════════════════════════
# UTILS
# ═══════════════════════════════════════════════════════════════════════════

def get_ordinal(n: int) -> str:
    """Returns 1st, 2nd, 3rd, etc."""
    if 11 <= (n % 100) <= 13:
        suffix = 'th'
    else:
        suffix = ['th', 'st', 'nd', 'rd', 'th', 'th', 'th', 'th', 'th', 'th'][n % 10]
    return f"{n}{suffix}"



# ═══════════════════════════════════════════════════════════════════════════
# AUTHOR LAYOUT ENGINE (IEEE)
# ═══════════════════════════════════════════════════════════════════════════

class IEEEAuthorLayoutManager:
    """
    Dynamically generates IEEE-compliant author grids for ReportLab PDFs.
    Handles 1-6 authors with specific row/column logic.
    """

    def __init__(self, styles: dict):
        self.styles = styles

    def _create_author_cell(self, author: AuthorInfo) -> List[Paragraph]:
        """Creates a list of paragraphs for a single author cell."""
        from reportlab.platypus import Paragraph
        
        lines = []
        style = self.styles.get('author_block', self.styles.get('body'))
        
        text = ""
        if author.name: text += f"<i>{html.escape(author.name)}</i><br/>"
        if author.role: text += f"<i>{html.escape(author.role)}</i><br/>"
        if author.department: text += f"<i>{html.escape(author.department)}</i><br/>"
        if author.institution: text += f"<i>{html.escape(author.institution)}</i><br/>"
        if author.university: text += f"<i>{html.escape(author.university)}</i><br/>"
        
        if author.address:
            addr_text = author.address
            if not addr_text.lower().startswith("address:"):
                addr_text = f"Address: {addr_text}"
            text += f"<i>{html.escape(addr_text)}</i><br/>"
            
        if author.pincode:
            pin = author.pincode
            if "pin code:" not in pin.lower() and "pincode:" not in pin.lower():
                pin = f"Pin code: {pin}"
            text += f"<i>{html.escape(pin)}</i><br/>"
            
        if author.email:
            # Add Mail ID label on one line, then email wrapped in blue, underlined, italic HTML tags on next line
            text += f"<i>Mail ID:</i><br/><u><font color='blue'><i>{html.escape(author.email)}</i></font></u>"
            
        if text.endswith("<br/>"): text = text[:-5]
        
        lines.append(Paragraph(text, style))
        return lines

    def generate_author_grid(self, authors: List[AuthorInfo], available_width: float) -> List:
        """Determines the grid structure (rows/cols) and returns a list of Table flowables."""
        from reportlab.platypus import Table, TableStyle, Paragraph
        
        num_authors = len(authors)
        if num_authors == 0:
            return []

        flowables = []
        data = []
        row_data = []
        
        col_w = available_width / 3.0

        for i, a in enumerate(authors):
            cell_content = self._create_author_cell(a)
            row_data.append(cell_content)
            
            if (i + 1) % 3 == 0:
                data.append(row_data)
                row_data = []
                
        if row_data:
            while len(row_data) < 3:
                row_data.append([Paragraph("", self.styles.get('body'))])
            data.append(row_data)
            
        t = Table(data, colWidths=[col_w]*3)
        t.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0,0), (-1,-1), 10),
            ('RIGHTPADDING', (0,0), (-1,-1), 10),
            ('TOPPADDING', (0,0), (-1,-1), 0),
            ('BOTTOMPADDING', (0,0), (-1,-1), 12),
        ]))
        
        return [t]


# ═══════════════════════════════════════════════════════════════════════════
# PDF GENERATOR  (ReportLab BaseDocTemplate)
# ═══════════════════════════════════════════════════════════════════════════

class PDFGenerator:
    """
    Generates IEEE-formatted PDF documents using ReportLab BaseDocTemplate
    with proper single-column and dual-column PageTemplate objects.
    """

    PAGE_W, PAGE_H = A4
    LEFT   = LEFT_MARGIN
    RIGHT  = RIGHT_MARGIN
    TOP    = TOP_MARGIN
    BOTTOM = BOTTOM_MARGIN

    def __init__(self, output_folder: str, style_config: dict = None):
        self.output_folder = output_folder
        self.style_config = style_config or {}
        self.styles = self._build_styles()

    # ─── Style Sheet ──────────────────────────────────────────────────────

    def _build_styles(self) -> dict:
        s = {}
        
        # Style overrides
        title_size = self.style_config.get('titleSize', 24)
        section_size = self.style_config.get('sectionSize', 10)
        subheading_size = self.style_config.get('subheadingSize', 10)
        body_size = self.style_config.get('bodySize', 10)
        line_spacing = self.style_config.get('lineSpacing', 1.0)
        font_family = self.style_config.get('fontFamily', 'Times-Roman')
        
        # Map Font Families
        rl_font = 'Times-Roman'
        rl_bold = 'Times-Bold'
        rl_italic = 'Times-Italic'
        rl_bold_italic = 'Times-BoldItalic'
        
        if font_family.lower() in ['arial', 'helvetica']:
            rl_font = 'Helvetica' # RL default for sans-serif
            rl_bold = 'Helvetica-Bold'
            rl_italic = 'Helvetica-Oblique'
            rl_bold_italic = 'Helvetica-BoldOblique'

        s['title'] = ParagraphStyle(
            'IEEE_Title',
            fontName=rl_bold,
            fontSize=24,
            leading=24 * 1.2 * line_spacing,
            alignment=TA_CENTER,
            spaceAfter=12,
        )

        s['author_name_italic'] = ParagraphStyle(
            'IEEE_AuthorNameItalic',
            fontName=rl_italic,
            fontSize=10,
            leading=12,
            alignment=TA_CENTER,
            spaceAfter=2,
        )
        s['author_aff_italic'] = ParagraphStyle(
            'IEEE_AuthorAffItalic',
            fontName=rl_italic,
            fontSize=10,
            leading=12,
            alignment=TA_CENTER,
            spaceAfter=2,
        )
        s['author_email_italic'] = ParagraphStyle(
            'IEEE_AuthorEmailItalic',
            fontName=rl_italic,
            fontSize=10,
            leading=12,
            alignment=TA_CENTER,
            spaceAfter=0,
        )
        s['author_block'] = ParagraphStyle(
            'IEEE_AuthorBlock',
            fontName='Times-Roman',
            fontSize=10,
            leading=12,
            alignment=TA_CENTER,
            spaceAfter=0,
        )
        s['abstract_heading'] = ParagraphStyle(
            'IEEE_AbstractHeading',
            fontName=rl_bold_italic,
            fontSize=9, # IEEE often uses smaller font (9pt) for Abstract
            leading=11,
            alignment=TA_JUSTIFY,
            spaceBefore=12,
            spaceAfter=2,
        )
        s['abstract_body'] = ParagraphStyle(
            'IEEE_AbstractBody',
            fontName=rl_bold_italic,
            fontSize=9,
            leading=11,
            alignment=TA_JUSTIFY,
        )
        s['keywords_body'] = ParagraphStyle(
            'IEEE_KeywordsBody',
            fontName=rl_bold_italic,
            fontSize=9,
            leading=11,
            alignment=TA_JUSTIFY,
            spaceAfter=12,
        )
        s['section_heading'] = ParagraphStyle(
            'IEEE_SectionHeading',
            fontName=rl_font,
            fontSize=section_size,
            leading=section_size * 1.2 * line_spacing,
            alignment=TA_CENTER,
            spaceBefore=12,
            spaceAfter=6,
            textTransform='uppercase', # We'll handle Small Caps via manual font switching if needed or just Uppercase
        )
        s['subsection_heading'] = ParagraphStyle(
            'IEEE_SubsectionHeading',
            fontName='Times-Italic',
            fontSize=10,
            leading=12,
            alignment=TA_LEFT,
            spaceBefore=9,
            spaceAfter=3,
        )
        s['body'] = ParagraphStyle(
            'IEEE_Body',
            fontName='Times-Roman',
            fontSize=10,
            leading=12,
            alignment=TA_JUSTIFY,
            firstLineIndent=0.422 * CM,
            spaceBefore=0,
            spaceAfter=0,
        )
        s['reference'] = ParagraphStyle(
            'IEEE_Reference',
            fontName='Times-Roman',
            fontSize=9,
            leading=11,
            alignment=TA_LEFT,
            leftIndent=14,
            firstLineIndent=-14,
            spaceAfter=3,
        )
        s['figure_caption'] = ParagraphStyle(
            'IEEE_FigureCaption',
            fontName='Times-Roman',
            fontSize=10,
            leading=12,
            alignment=TA_CENTER,
            spaceBefore=6,
            spaceAfter=6,
            textTransform='uppercase', # Scopus/IEEE Small Caps for Fig. X.
        )
        s['table_caption'] = ParagraphStyle(
            'IEEE_TableCaption',
            fontName='Times-Roman',
            fontSize=10,
            leading=12,
            alignment=TA_CENTER,
            spaceBefore=6,
            spaceAfter=6,
            textTransform='uppercase',
        )
        s['list_item'] = ParagraphStyle(
            'IEEE_ListItem',
            fontName='Times-Roman',
            fontSize=10,
            leading=12,
            alignment=TA_JUSTIFY,
            leftIndent=0.844 * CM, # 0.422 * 2
            firstLineIndent=-0.422 * CM,
            spaceBefore=0,
            spaceAfter=0,
        )
        return s

    # ─── Page Layout ──────────────────────────────────────────────────────

    def _build_doc(self, filepath: str) -> BaseDocTemplate:
        doc = BaseDocTemplate(
            filepath,
            pagesize=A4,
            leftMargin=self.LEFT,
            rightMargin=self.RIGHT,
            topMargin=self.TOP,
            bottomMargin=self.BOTTOM,
        )

        usable_w = self.PAGE_W - self.LEFT - self.RIGHT
        usable_h = self.PAGE_H - self.TOP - self.BOTTOM

        # ── Single-column template ──
        frame_1col = Frame(
            self.LEFT, self.BOTTOM,
            usable_w, usable_h,
            id='single',
            leftPadding=0, rightPadding=0,
            topPadding=0, bottomPadding=0,
        )
        tmpl_1col = PageTemplate(id='OneCol', frames=[frame_1col])

        # ── Two-column template ──
        # IEEE Specs: Column Width = 8.25 cm, Inter Column Space = 0.8 cm
        col_w = COL_W
        col_gap = COL_GAP
        
        frame_left = Frame(
            self.LEFT, self.BOTTOM,
            col_w, usable_h,
            id='left',
            leftPadding=0, rightPadding=0,
            topPadding=0, bottomPadding=0,
        )
        frame_right = Frame(
            self.LEFT + col_w + col_gap, self.BOTTOM,
            col_w, usable_h,
            id='right',
            leftPadding=0, rightPadding=0,
            topPadding=0, bottomPadding=0,
        )
        tmpl_2col = PageTemplate(id='TwoCol', frames=[frame_left, frame_right])

        doc.addPageTemplates([tmpl_1col, tmpl_2col])
        return doc

    # ─── Builder Methods ──────────────────────────────────────────────────

    def _build_title(self, story: list, data: DocumentData):
        if data.title:
            # IEEE Title: Capitalize Nouns, Pronouns, Verbs, Adjectives, Adverbs
            ignore_words = {'a', 'an', 'the', 'and', 'but', 'or', 'for', 'nor', 'on', 'at', 'to', 'from', 'by', 'of', 'in', 'with'}
            words = data.title.split()
            formatted_words = []
            for i, word in enumerate(words):
                clean_word = word.lower()
                if clean_word not in ignore_words or i == 0 or i == len(words) - 1:
                    formatted_words.append(word.capitalize())
                else:
                    formatted_words.append(clean_word)
            title_text = html.escape(" ".join(formatted_words))
            
            story.append(Spacer(1, 1.0 * inch - self.TOP))
            story.append(Paragraph(title_text, self.styles['title']))
            story.append(Spacer(1, 12))

    def _build_authors(self, story: list, data: DocumentData):
        if not data.authors:
            return

        available_width = self.PAGE_W - self.LEFT - self.RIGHT
        manager = IEEEAuthorLayoutManager(self.styles)
        
        story.append(Spacer(1, 12))
        grid_flowables = manager.generate_author_grid(data.authors, available_width)
        story.extend(grid_flowables)
        story.append(Spacer(1, 18))

        story.append(Spacer(1, 18)) # Space after Author Section: 18 pt

        story.append(Spacer(1, 18)) # Space after Author Section: 18 pt


    def _build_abstract(self, story: list, data: DocumentData):
        """Finds Abstract in sections, removes it, and renders it here."""
        abstract_sec = None
        for i, sec in enumerate(data.sections):
            if sec.heading.upper() == "ABSTRACT":
                abstract_sec = data.sections.pop(i)
                break
        
        if abstract_sec and abstract_sec.body:
            # IEEE Style: Abstract— inline
            text = f"<b><i>Abstract—</i></b>{html.escape(abstract_sec.body)}"
            story.append(Paragraph(text, self.styles['abstract_body']))

    def _build_keywords(self, story: list, data: DocumentData):
        if data.keywords:
            # Sort alphabetically, only first keyword capitalized
            keywords = sorted([k.strip() for k in data.keywords])
            if keywords:
                keywords[0] = keywords[0].capitalize()
                for i in range(1, len(keywords)):
                    keywords[i] = keywords[i].lower()
                
                # IEEE Style: Index Terms— inline
                kw_text = f"<b><i>Index Terms—</i></b>{html.escape(', '.join(keywords))}"
                story.append(Paragraph(kw_text, self.styles['keywords_body']))

    def _switch_to_two_column(self, story: list):
        """Trigger 2-column layout. Uses FrameBreak to start flowing into left column."""
        story.append(NextPageTemplate('TwoCol'))
        story.append(FrameBreak())

    def _build_sections(self, story: list, data: DocumentData):
        for sec in data.sections:
            # IEEE Primary Heading: Small Caps (simulated with upper() and 10pt)
            story.append(Paragraph(html.escape(sec.heading.upper()), self.styles['section_heading']))

            if sec.body:
                self._render_body_with_media(story, sec, self.styles['body'])

            story.append(Spacer(1, 6))

    def _render_body_with_media(self, story: list, sec: SectionData, body_style):
        """Helper to render body text while handling inline media placeholders."""
        # Split body into parts, handling placeholders
        parts = re.split(r'(\[\[FIG:\d+\]\]|\[\[TBL:\d+\]\]|\[\[EQ:\d+\]\])', sec.body)
        
        for part in parts:
            if not part.strip():
                continue

            # Handle placeholders
            fig_match = re.match(r'\[\[FIG:(\d+)\]\]', part)
            tbl_match = re.match(r'\[\[TBL:(\d+)\]\]', part)
            eq_match = re.match(r'\[\[EQ:(\d+)\]\]', part)

            if fig_match:
                idx = int(fig_match.group(1))
                if idx < len(sec.figures):
                    fig = sec.figures[idx]
                    self._add_figure(story, fig['path'], fig['caption'])
            elif tbl_match:
                idx = int(tbl_match.group(1))
                if idx < len(sec.tables):
                    tbl = sec.tables[idx]
                    self._add_table(story, tbl['data'], tbl['caption'])
            elif eq_match:
                idx = int(eq_match.group(1))
                if idx < len(sec.equations):
                    eq = sec.equations[idx]
                    self._add_equation(story, eq['text'], eq['num'])
            else:
                # Regular paragraph text
                paragraphs = part.split('\n')
                for para_text in paragraphs:
                    stripped = para_text.strip()
                    if not stripped:
                        continue
                    
                    # Matches I., A., 1.1, etc.
                    if re.match(r'^([A-Z]\.|[0-9]+\.[0-9]+|[IVX]+\.)\s+\S', stripped):
                        story.append(Paragraph(html.escape(stripped), self.styles['subsection_heading']))
                    elif stripped.startswith('[BULLET]'):
                        # Handle Bullet: Replace marker with a bullet character
                        content = stripped[8:].strip()
                        story.append(Paragraph(f"• {html.escape(content)}", self.styles['list_item']))
                    else:
                        story.append(Paragraph(html.escape(stripped), body_style))

    def _build_references(self, story: list, data: DocumentData):
        if not data.references:
            return

        story.append(Paragraph("REFERENCES", self.styles['section_heading']))
        for i, ref in enumerate(data.references):
            # Ensure sequence [1]
            ref_text = ref.strip()
            if not ref_text.startswith('['):
                ref_text = f"[{i+1}] {ref_text}"
            story.append(Paragraph(html.escape(ref_text), self.styles['reference']))

    def _add_figure(self, story: list, img_path: str, caption: str):
        from reportlab.platypus import Image
        if not os.path.exists(img_path):
            return
        img = Image(img_path)
        # Scale to max column width
        if img.drawWidth > COL_W:
            ratio = COL_W / img.drawWidth
            img.drawWidth = COL_W
            img.drawHeight *= ratio
        story.append(img)
        story.append(Paragraph(html.escape(caption), self.styles['figure_caption']))

    def _add_table(self, story: list, table_data: list, caption: str):
        from reportlab.platypus import Table, TableStyle
        story.append(Paragraph(html.escape(caption), self.styles['table_caption']))
        
        if not table_data:
            return
            
        # Ensure all rows have the same number of columns
        max_cols = max(len(row) for row in table_data)
        padded_data = []
        for row in table_data:
            pad = [''] * (max_cols - len(row))
            padded_data.append(row + pad)
            
        # Escape table cells
        esc_data = [[html.escape(str(cell)) for cell in row] for row in padded_data]
        tbl = Table(esc_data)
        tbl.setStyle(TableStyle([
            ('FONT', (0, 0), (-1, 0), 'Times-Bold'),
            ('FONT', (0, 1), (-1, -1), 'Times-Roman'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ]))
        story.append(tbl)

    def _add_equation(self, story: list, eq_text: str, num: int):
        from reportlab.platypus import Table, TableStyle
        # IEEE Equation: Centered equation with right-aligned number (n)
        # Use a table with 3 columns: empty, content, number
        data = [['', Paragraph(f"<i>{html.escape(eq_text)}</i>", self.styles['body']), f"({num})"]]
        tbl = Table(data, colWidths=[20, COL_W - 60, 40])
        tbl.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('ALIGN', (1, 0), (1, 0), 'CENTER'),
            ('ALIGN', (2, 0), (2, 0), 'RIGHT'),
        ]))
        story.append(Spacer(1, 6))
        story.append(tbl)
        story.append(Spacer(1, 6))

    # ─── Main Entry Point ─────────────────────────────────────────────────

    def generate_pdf(self, data: DocumentData, filename: str) -> str:
        """
        Builds and saves an IEEE-formatted PDF. Returns the output file path.
        """
        import re
        filepath = os.path.join(self.output_folder, filename)
        doc = self._build_doc(filepath)

        story = []

        # ── Single-column block: Title, Authors, Abstract, Keywords ──
        self._build_title(story, data)
        self._build_authors(story, data)

        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.black, spaceAfter=12))

        self._build_abstract(story, data)
        self._build_keywords(story, data)

        # ── Switch to 2-column ──
        self._switch_to_two_column(story)

        # ── Two-column block: Sections, References ──
        self._build_sections(story, data)
        self._build_references(story, data)

        # Explicitly remove any default page numbers (BaseDocTemplate doesn't add them by default)
        # but we can ensure it by not defining any onPage handlers.
        
        doc.build(story)
        return filepath


# ═══════════════════════════════════════════════════════════════════════════
# DOCX GENERATOR  (python-docx)
# ═══════════════════════════════════════════════════════════════════════════

class WordGenerator:
    """
    Generates IEEE-formatted DOCX documents using python-docx.
    Uses a continuous section break with XML-level w:cols="2" for two-column layout.
    """

    def __init__(self, output_folder: str, style_config: dict = None):
        self.output_folder = output_folder
        self.style_config = style_config or {}

    # ─── Paragraph Helpers ────────────────────────────────────────────────

    def _add_paragraph(self, doc: Document, text: str,
                       bold=False, italic=False, font_size=None,
                       alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                       space_before=0, space_after=0,
                       first_line_indent=0, left_indent=0) -> None:
        
        # Style overrides
        f_size = font_size or self.style_config.get('bodySize', 10)
        f_name = self.style_config.get('fontFamily', 'Times New Roman')
        line_spacing = self.style_config.get('lineSpacing', 1.0)

        p = doc.add_paragraph()
        p.alignment = alignment
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        p.paragraph_format.line_spacing = line_spacing if line_spacing > 0 else 1.0
        
        if first_line_indent:
            p.paragraph_format.first_line_indent = first_line_indent
        if left_indent:
            p.paragraph_format.left_indent = left_indent

        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.name = f_name
        run.font.size = Pt(f_size)
        return p

    def _add_section_heading(self, doc: Document, text: str) -> None:
        section_size = self.style_config.get('sectionSize', 10)
        f_name = self.style_config.get('fontFamily', 'Times New Roman')
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(6)
        run = p.add_run(text.upper())
        run.font.name = f_name
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)

    def _add_subsection_heading(self, doc: Document, text: str) -> None:
        subheading_size = self.style_config.get('subheadingSize', 10)
        f_name = self.style_config.get('fontFamily', 'Times New Roman')
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(9)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(text)
        run.italic = True
        run.font.name = f_name
        run.font.size = Pt(subheading_size)

    # ─── Column Switch ────────────────────────────────────────────────────

    def _switch_to_two_column(self, doc: Document) -> None:
        """
        Inserts a continuous section break and sets the new section to 2 columns
        via direct XML manipulation (w:cols w:num="2" w:space="720").
        """
        new_section = doc.add_section(WD_SECTION.CONTINUOUS)
        sect_pr = new_section._sectPr

        # Remove any existing w:cols elements
        for existing in sect_pr.findall(qn('w:cols')):
            sect_pr.remove(existing)

        cols = OxmlElement('w:cols')
        cols.set(qn('w:num'), '2')
        cols.set(qn('w:space'), '454')   # ~0.8cm in twentieths of a point (0.8 / 2.54 * 72 * 20 = 453.5)
        cols.set(qn('w:equalWidth'), '1')
        sect_pr.append(cols)

    # ─── Builder Methods ──────────────────────────────────────────────────

    def _build_title(self, doc: Document, data: DocumentData) -> None:
        if not data.title:
            return
            
        title_size = self.style_config.get('titleSize', 24)
        f_name = self.style_config.get('fontFamily', 'Times New Roman')
        
        # IEEE Title Capitalization logic
        ignore_words = {'a', 'an', 'the', 'and', 'but', 'or', 'for', 'nor', 'on', 'at', 'to', 'from', 'by', 'of', 'in', 'with'}
        words = data.title.split()
        formatted_words = []
        for i, word in enumerate(words):
            clean_word = word.lower()
            if clean_word not in ignore_words or i == 0 or i == len(words) - 1:
                formatted_words.append(word.capitalize())
            else:
                formatted_words.append(clean_word)
        title_text = " ".join(formatted_words)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(12) # Space After Title: 12 pt
        
        # IEEE Title: Capitalize Nouns, Pronouns, Verbs, Adjectives, Adverbs
        ignore_words = {'a', 'an', 'the', 'and', 'but', 'or', 'for', 'nor', 'on', 'at', 'to', 'from', 'by', 'of', 'in', 'with'}
        words = data.title.split()
        formatted_words = []
        for i, word in enumerate(words):
            clean_word = word.lower()
            if clean_word not in ignore_words or i == 0 or i == len(words) - 1:
                formatted_words.append(word.capitalize())
            else:
                formatted_words.append(clean_word)
        title_text = " ".join(formatted_words)

        run = p.add_run(title_text)
        run.bold = True
        run.font.name = f_name
        run.font.size = Pt(24) # Title Size: 24 pt

    def _build_authors(self, doc: Document, data: DocumentData):
        if not data.authors:
            return

        # Space before Author Section: 12 pt
        doc.add_paragraph().paragraph_format.space_after = Pt(12)

        num_authors = len(data.authors)
        import math
        rows_needed = math.ceil(num_authors / 3)
        if rows_needed == 0:
            rows_needed = 1
        table = doc.add_table(rows=rows_needed, cols=3)
        
        from docx.oxml.shared import OxmlElement
        from docx.oxml.ns import qn
        tbl_pr = table._element.xpath('w:tblPr')
        if tbl_pr:
            jc = OxmlElement('w:jc')
            jc.set(qn('w:val'), 'center')
            tbl_pr[0].append(jc)
            
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        total_width = 6.5
        col_w = total_width / 3
        
        for row in table.rows:
            for cell in row.cells:
                cell.width = Inches(col_w)
                
        author_index = 0
        for row in range(rows_needed):
            for col in range(3):
                cell = table.rows[row].cells[col]
                if author_index < num_authors:
                    a = data.authors[author_index]
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.0
                    
                    if a.name:
                        run = p.add_run(a.name + "\n")
                        run.italic = True
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(10)
                        
                    def add_text(text, new_line=True):
                        r = p.add_run(text + ("\n" if new_line else ""))
                        r.italic = True
                        r.font.name = 'Times New Roman'
                        r.font.size = Pt(10)

                    if a.role: add_text(a.role)
                    if a.department: add_text(a.department)
                    if a.institution: add_text(a.institution)
                    if a.university: add_text(a.university)
                    
                    if a.address:
                        addr_text = a.address
                        if not addr_text.lower().startswith("address:"):
                            addr_text = f"Address: {addr_text}"
                        add_text(addr_text)
                        
                    if a.pincode:
                        pin = a.pincode
                        if "pin code:" not in pin.lower() and "pincode:" not in pin.lower():
                            pin = f"Pin code: {pin}"
                        add_text(pin)
                        
                    if a.email:
                        add_text("Mail ID: ")
                        # Format the email on a new line, italic, underlined, and blue
                        r_email = p.add_run(a.email)
                        r_email.italic = True
                        r_email.underline = True
                        r_email.font.color.rgb = RGBColor(0, 0, 255)
                        r_email.font.name = 'Times New Roman'
                        r_email.font.size = Pt(10)
                        
                    # Remove trailing newline if it's the last element and has newline
                    if len(p.runs) > 0 and p.runs[-1].text.endswith("\n"):
                        p.runs[-1].text = p.runs[-1].text[:-1]

                    author_index += 1

        # Space after Author Section: 18 pt
        doc.add_paragraph().paragraph_format.space_after = Pt(18)


    def _build_abstract(self, doc: Document, data: DocumentData):
        """Finds Abstract in sections, removes it, and renders it here."""
        abstract_sec = None
        for i, sec in enumerate(data.sections):
            if sec.heading.upper() == "ABSTRACT":
                abstract_sec = data.sections.pop(i)
                break
        
        if abstract_sec and abstract_sec.body:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(2)
            
            # IEEE Style: Abstract— inline (Bold Italic)
            run_label = p.add_run("Abstract—")
            run_label.font.name = 'Times New Roman'
            run_label.font.size = Pt(9)
            run_label.bold = True
            run_label.italic = True
            
            run_body = p.add_run(abstract_sec.body)
            run_body.font.name = 'Times New Roman'
            run_body.font.size = Pt(9)
            run_body.bold = True
            run_body.italic = True

    def _build_keywords(self, doc: Document, data: DocumentData) -> None:
        if not data.keywords:
            return

        # Sort alphabetically, only first keyword capitalized
        keywords = sorted([k.strip() for k in data.keywords])
        if keywords:
            keywords[0] = keywords[0].capitalize()
            for i in range(1, len(keywords)):
                keywords[i] = keywords[i].lower()
            
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(12)
            
            # IEEE Style: Index Terms— inline
            run_label = p.add_run("Index Terms—")
            run_label.font.name = 'Times New Roman'
            run_label.font.size = Pt(9)
            run_label.bold = True
            run_label.italic = True
            
            run_body = p.add_run(", ".join(keywords))
            run_body.font.name = 'Times New Roman'
            run_body.font.size = Pt(9)
            run_body.bold = True
            run_body.italic = True

    def _build_sections(self, doc: Document, data: DocumentData) -> None:
        for sec in data.sections:
            self._add_section_heading(doc, sec.heading)

            if sec.body:
                self._render_body_with_media(doc, sec)

    def _render_body_with_media(self, doc: Document, sec: SectionData, 
                               bold=False, italic=False, font_size=10):
        """Helper to render body text in Word while handling inline media placeholders."""
        paragraphs = re.split(r'(\[\[FIG:\d+\]\]|\[\[TBL:\d+\]\]|\[\[EQ:\d+\]\])', sec.body)
        
        for para_part in paragraphs:
            if not para_part.strip():
                continue

            # Handle placeholders
            fig_match = re.match(r'\[\[FIG:(\d+)\]\]', para_part)
            tbl_match = re.match(r'\[\[TBL:(\d+)\]\]', para_part)
            eq_match = re.match(r'\[\[EQ:(\d+)\]\]', para_part)

            if fig_match:
                idx = int(fig_match.group(1))
                if idx < len(sec.figures):
                    fig = sec.figures[idx]
                    self._add_figure(doc, fig['path'], fig['caption'])
            elif tbl_match:
                idx = int(tbl_match.group(1))
                if idx < len(sec.tables):
                    tbl = sec.tables[idx]
                    self._add_table(doc, tbl['data'], tbl['caption'])
            elif eq_match:
                idx = int(eq_match.group(1))
                if idx < len(sec.equations):
                    eq = sec.equations[idx]
                    self._add_equation(doc, eq['text'], eq['num'])
            else:
                # Regular paragraph text
                lines = para_part.split('\n')
                for para_text in lines:
                    stripped = para_text.strip()
                    if not stripped:
                        continue
                    
                    if re.match(r'^([A-Z]\.|[0-9]+\.[0-9]+|[IVX]+\.)\s+\S', stripped):
                        self._add_subsection_heading(doc, stripped)
                    elif stripped.startswith('[BULLET]'):
                        content = stripped[8:].strip()
                        try:
                            p = doc.add_paragraph(content, style='List Bullet')
                        except:
                            self._add_paragraph(doc, f"• {content}", font_size=font_size,
                                               bold=bold, italic=italic)
                    else:
                        p = self._add_paragraph(doc, stripped, font_size=font_size, 
                                            bold=bold, italic=italic)

    def _build_references(self, doc: Document, data: DocumentData) -> None:
        if not data.references:
            return

        self._add_section_heading(doc, "References")

        for i, ref in enumerate(data.references):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent   = Pt(14)
            p.paragraph_format.first_line_indent = Pt(-14)
            p.paragraph_format.space_after   = Pt(2)
            
            ref_text = ref.strip()
            if not ref_text.startswith('['):
                ref_text = f"[{i+1}] {ref_text}"
            
            run = p.add_run(ref_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(9)

    def _add_figure(self, doc: Document, img_path: str, caption: str):
        from docx.shared import Cm
        if not os.path.exists(img_path):
            return
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Cm(8.25))
        
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(6)
        cap.paragraph_format.space_after = Pt(6)
        run_cap = cap.add_run(caption.upper())
        run_cap.font.name = 'Times New Roman'
        run_cap.font.size = Pt(10)
        run_cap.font.small_caps = True

    def _add_table(self, doc: Document, table_data: list, caption: str):
        if not table_data:
            return
            
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(6)
        cap.paragraph_format.space_after = Pt(6)
        run_cap = cap.add_run(caption.upper())
        run_cap.font.name = 'Times New Roman'
        run_cap.font.size = Pt(10)
        run_cap.font.small_caps = True

        max_cols = max(len(row) for row in table_data)
        table = doc.add_table(rows=len(table_data), cols=max_cols)
        table.style = 'Table Grid'
        
        for i, row in enumerate(table_data):
            for j in range(max_cols):
                val = str(row[j]) if j < len(row) else ''
                table.cell(i, j).text = val

    def _add_equation(self, doc: Document, eq_text: str, num: int):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(eq_text)
        run.italic = True
        tab_run = p.add_run('\t')
        num_run = p.add_run(f"({num})")
        # Ensure right alignment for number via tab stops (would need section width)
        # For now, just adding them.

    # ─── Main Entry Point ─────────────────────────────────────────────────

    def generate_docx(self, data: DocumentData, filename: str) -> str:
        """
        Builds and saves an IEEE-formatted DOCX. Returns the output file path.
        """
        filepath = os.path.join(self.output_folder, filename)
        doc = Document()

        # ── Global default font ──
        doc.styles['Normal'].font.name = 'Times New Roman'
        doc.styles['Normal'].font.size = Pt(10)

        # ── Page margins (IEEE A4) ──
        for section in doc.sections:
            section.top_margin    = Cm(2.54)
            section.bottom_margin = Cm(2.86)
            section.left_margin   = Cm(1.9)
            section.right_margin  = Cm(1.9)
            section.header_distance = Pt(0)
            section.footer_distance = Pt(0)

        # ── Single-column block ──
        self._build_title(doc, data)
        self._build_authors(doc, data)
        
        # Horizontal Line (simulated with a thin-bordered table or blank line)
        doc.add_paragraph().paragraph_format.space_after = Pt(12)

        self._build_abstract(doc, data)
        self._build_keywords(doc, data)

        # ── Switch to two columns ──
        self._switch_to_two_column(doc)

        # ── Two-column block ──
        self._build_sections(doc, data)
        self._build_references(doc, data)

        doc.save(filepath)
        return filepath


# ═══════════════════════════════════════════════════════════════════════════
# SPRINGER PDF GENERATOR
# ═══════════════════════════════════════════════════════════════════════════

class SpringerPDFGenerator(PDFGenerator):
    """
    Generates Springer-formatted PDF documents.
    Single-column, specific font sizes, and mandatory section handling.
    """

    def _build_styles(self) -> dict:
        s = super()._build_styles()
        
        line_spacing = self.style_config.get('lineSpacing', 1.15)
        
        # Springer Specific overrides
        s['title'] = ParagraphStyle(
            'Springer_Title',
            fontName='Times-Roman',
            fontSize=14,
            leading=14 * 1.2 * line_spacing,
            alignment=TA_CENTER,
            spaceAfter=24, # Spacing Section: 24pt after title
        )
        s['author_block'] = ParagraphStyle(
            'Springer_Author',
            fontName='Times-Roman',
            fontSize=10,
            leading=12,
            alignment=TA_CENTER,
            spaceAfter=6,
        )
        s['affiliation'] = ParagraphStyle(
            'Springer_Affiliation',
            fontName='Times-Roman',
            fontSize=9,
            leading=11,
            alignment=TA_CENTER,
            spaceAfter=2,
        )
        s['abstract_heading'] = ParagraphStyle(
            'Springer_AbstractHeading',
            fontName='Times-Bold',
            fontSize=12,
            leading=14,
            alignment=TA_LEFT,
            spaceBefore=12,
            spaceAfter=6,
        )
        s['abstract_body'] = ParagraphStyle(
            'Springer_AbstractBody',
            fontName='Times-Roman',
            fontSize=10,
            leading=10 * line_spacing,
            alignment=TA_JUSTIFY,
        )
        s['keywords_label'] = ParagraphStyle(
            'Springer_KeywordsLabel',
            fontName='Times-Italic',
            fontSize=10,
            leading=12,
            alignment=TA_LEFT,
            spaceBefore=6,
        )
        s['section_heading'] = ParagraphStyle(
            'Springer_SectionHeading',
            fontName='Times-Bold',
            fontSize=10,
            leading=12,
            alignment=TA_LEFT,
            spaceBefore=12,
            spaceAfter=6,
        )
        s['body'] = ParagraphStyle(
            'Springer_Body',
            fontName='Times-Roman',
            fontSize=10,
            leading=10 * line_spacing,
            alignment=TA_JUSTIFY,
            spaceBefore=0,
            spaceAfter=0,
        )
        return s

    def _build_authors_springer(self, story: list, data: DocumentData):
        if not data.authors:
            return

        # Format Author Line: Author1,2*, Author2,3†...
        author_parts = []
        for a in data.authors:
            part = a.name
            suffixes = [str(sid) for sid in a.aff_ids]
            if a.is_corresponding: suffixes.append("*")
            if a.equal_contrib: suffixes.append("†")
            
            if suffixes:
                part += f"<sup>{','.join(suffixes)}</sup>"
            author_parts.append(part)
        
        author_text = ", ".join(author_parts)
        if len(author_parts) > 1:
            # Replace last comma with " and "
            last_comma_idx = author_text.rfind(", ")
            if last_comma_idx != -1:
                author_text = author_text[:last_comma_idx] + " and " + author_text[last_comma_idx+2:]

        story.append(Paragraph(author_text, self.styles['author_block']))

        # Spacing Rule: Authors -> Affiliations = 6pt
        story.append(Spacer(1, 6))

        # Affiliations: Font 9pt, Regular
        # Rule: If affiliation ID belongs to a corresponding author, append *
        corr_aff_ids = set()
        for a in data.authors:
            if a.is_corresponding:
                corr_aff_ids.update(a.aff_ids)

        for idx, aff in data.affiliations.items():
            lab = str(idx)
            if idx in corr_aff_ids:
                lab += "*"
            story.append(Paragraph(f"<sup>{lab}</sup>{html.escape(aff)}", self.styles['affiliation']))

        # Spacing Rule: Affiliations -> Email = 12pt
        story.append(Spacer(1, 12))

        # Corresponding Author E-mails
        # Rule: Use semicolon between emails AND at the end. Blue/Underlined.
        corr_emails = [a.email for a in data.authors if a.is_corresponding and a.email]
        if corr_emails:
            formatted_emails = [f'<font color="blue"><u>{html.escape(e)}</u></font>' for e in corr_emails]
            email_text = f"*Corresponding author(s). E-mail(s): {'; '.join(formatted_emails)};"
            story.append(Paragraph(email_text, self.styles['affiliation']))
        
        # Contributing authors (others)
        other_emails = [a.email for a in data.authors if not a.is_corresponding and a.email]
        if other_emails:
            formatted_emails = [f'<font color="blue"><u>{html.escape(e)}</u></font>' for e in other_emails]
            email_text = f"Contributing authors: {'; '.join(formatted_emails)};"
            story.append(Paragraph(email_text, self.styles['affiliation']))
        
        # Equal Contribution Statement
        # Rule: Email -> Equal Contribution = 6pt
        any_equal = any(a.equal_contrib for a in data.authors)
        if any_equal:
            story.append(Spacer(1, 6))
            story.append(Paragraph("†These authors contributed equally to this work.", self.styles['affiliation']))
        
        story.append(Spacer(1, 12))

    def generate_pdf(self, data: DocumentData, filename: str) -> str:
        filepath = os.path.join(self.output_folder, filename)
        doc = BaseDocTemplate(
            filepath,
            pagesize=A4,
            leftMargin=self.LEFT,
            rightMargin=self.RIGHT,
            topMargin=self.TOP,
            bottomMargin=self.BOTTOM,
        )

        # Springer is typically single column
        frame = Frame(self.LEFT, self.BOTTOM, self.PAGE_W - self.LEFT - self.RIGHT, self.PAGE_H - self.TOP - self.BOTTOM, id='normal')
        template = PageTemplate(id='Springer', frames=frame)
        doc.addPageTemplates([template])

        story = []
        self._build_title(story, data)
        self._build_authors_springer(story, data)
        
        # Abstract & Keywords
        for sec in data.sections:
            if sec.heading.upper() == "ABSTRACT":
                story.append(Paragraph("Abstract", self.styles['abstract_heading']))
                story.append(Paragraph(html.escape(sec.body), self.styles['abstract_body']))
                if data.keywords:
                    kw_text = f"<i>Keywords: {', '.join(data.keywords)}</i>"
                    story.append(Paragraph(kw_text, self.styles['keywords_label']))
                story.append(Spacer(1, 12))
                break

        # Sections (excluding Abstract)
        sec_num = 1
        for sec in data.sections:
            if sec.heading.upper() == "ABSTRACT": continue
            
            # Simple numbering
            h_text = f"{sec_num} {sec.heading}"
            story.append(Paragraph(html.escape(h_text), self.styles['section_heading']))
            if sec.body:
                self._render_body_with_media(story, sec, self.styles['body'])
            story.append(Spacer(1, 6))
            sec_num += 1

        self._build_references(story, data)
        doc.build(story)
        return filepath


# ═══════════════════════════════════════════════════════════════════════════
# SPRINGER WORD GENERATOR
# ═══════════════════════════════════════════════════════════════════════════

class SpringerWordGenerator(WordGenerator):
    """
    Generates Springer-formatted DOCX documents.
    """
    
    def _build_title(self, doc: Document, data: DocumentData) -> None:
        if not data.title: return
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(24)
        run = p.add_run(data.title)
        run.bold = False
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'

    def _build_authors_springer(self, doc: Document, data: DocumentData):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Space from Title to Authors = 12pt (Handled by title space_after)
        
        for i, a in enumerate(data.authors):
            run = p.add_run(a.name)
            run.font.size = Pt(10)
            run.bold = False # Rule: Do NOT bold author names
            
            # Superscripts
            suffixes = [str(sid) for sid in a.aff_ids]
            if a.is_corresponding: suffixes.append("*")
            if a.equal_contrib: suffixes.append("†")
            if suffixes:
                sup = p.add_run(",".join(suffixes))
                sup.font.superscript = True
                sup.font.size = Pt(10)
            
            if i < len(data.authors) - 2:
                p.add_run(", ")
            elif i == len(data.authors) - 2:
                p.add_run(" and ")
        
        p.paragraph_format.space_after = Pt(6) # Spacing: Authors -> Affiliations = 6pt

        # Affiliations: Font 9pt
        corr_aff_ids = set()
        for a in data.authors:
            if a.is_corresponding:
                corr_aff_ids.update(a.aff_ids)

        for idx, aff in data.affiliations.items():
            ap = doc.add_paragraph()
            ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            lab = str(idx)
            if idx in corr_aff_ids:
                lab += "*"
            
            run_idx = ap.add_run(lab)
            run_idx.font.superscript = True
            run_idx.font.size = Pt(9)
            
            run_aff = ap.add_run(f"{aff}")
            run_aff.font.size = Pt(9)
            ap.paragraph_format.space_after = Pt(2)
        
        last_aff_p = doc.paragraphs[-1]
        last_aff_p.paragraph_format.space_after = Pt(12)

        # Emails: Semicolon separator and trailing semicolon. Blue/Underlined.
        corr_emails = [a.email for a in data.authors if a.is_corresponding and a.email]
        if corr_emails:
            ep = doc.add_paragraph("*Corresponding author(s). E-mail(s): ")
            ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for i, email in enumerate(corr_emails):
                run = ep.add_run(email)
                run.font.color.rgb = RGBColor(0, 0, 255)
                run.underline = True
                if i < len(corr_emails) - 1:
                    ep.add_run("; ")
            ep.add_run(";")
            ep.paragraph_format.space_after = Pt(2)

        other_emails = [a.email for a in data.authors if not a.is_corresponding and a.email]
        if other_emails:
            ep = doc.add_paragraph("Contributing authors: ")
            ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for i, email in enumerate(other_emails):
                run = ep.add_run(email)
                run.font.color.rgb = RGBColor(0, 0, 255)
                run.underline = True
                if i < len(other_emails) - 1:
                    ep.add_run("; ")
            ep.add_run(";")
            ep.paragraph_format.space_after = Pt(2)

        # Equal Contribution
        any_equal = any(a.equal_contrib for a in data.authors)
        if any_equal:
            last_p = doc.paragraphs[-1]
            last_p.paragraph_format.space_after = Pt(6)
            eqp = doc.add_paragraph("†These authors contributed equally to this work.")
            eqp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            eqp.paragraph_format.space_after = Pt(12)

    def generate_docx(self, data: DocumentData, filename: str) -> str:
        filepath = os.path.join(self.output_folder, filename)
        doc = Document()
        doc.styles['Normal'].font.name = 'Times New Roman'
        doc.styles['Normal'].font.size = Pt(10)

        self._build_title(doc, data)
        self._build_authors_springer(doc, data)
        
        # Springer single column - no section break needed for columns
        
        # Abstract
        for sec in data.sections:
            if sec.heading.upper() == "ABSTRACT":
                h = doc.add_paragraph()
                run = h.add_run("Abstract")
                run.bold = True
                run.font.size = Pt(12)
                
                body = doc.add_paragraph(sec.body)
                body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                body.paragraph_format.line_spacing = 1.15
                
                if data.keywords:
                    kw = doc.add_paragraph()
                    run = kw.add_run(f"Keywords: {', '.join(data.keywords)}")
                    run.italic = True
                break
        
        # Sections
        sec_num = 1
        for sec in data.sections:
            if sec.heading.upper() == "ABSTRACT": continue
            h = doc.add_paragraph()
            run = h.add_run(f"{sec_num} {sec.heading}")
            run.bold = True
            
            if sec.body:
                self._render_body_with_media(doc, sec)
            sec_num += 1

        self._build_references(doc, data)
        doc.save(filepath)
        return filepath

# If old code passes a plain dict instead of DocumentData, wrap it.

def _dict_to_doc_data(sections_dict: dict) -> DocumentData:
    """Convert legacy flat dict (section_name → body) to DocumentData."""
    from nlp_processor import DocumentData, SectionData, AuthorInfo
    doc = DocumentData()
    doc.title = "Research Paper Output"
    for name, body in sections_dict.items():
        n_upper = name.upper()
        if 'ABSTRACT' in n_upper:
            doc.abstract = SectionData(heading="Abstract", body=body)
        elif 'KEYWORD' in n_upper:
            doc.keywords = [k.strip() for k in body.split(',') if k.strip()]
        elif 'REFERENCE' in n_upper:
            doc.references = [r.strip() for r in body.split('\n') if r.strip()]
        else:
            doc.sections.append(SectionData(heading=name, body=body))
    return doc


import re  # needed in _build_sections bodies at module level
