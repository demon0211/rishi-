from docx import Document
import os
import re

file_path = "uploads/CKD_ML_Prediction_Enhanced_Paper_111.docx"
if os.path.exists(file_path):
    doc = Document(file_path)
    found_any = False
    for i, p in enumerate(doc.paragraphs):
        xml = p._element.xml
        # Precise check for numbering properties in XML
        if 'w:numPr' in xml or 'List' in p.style.name:
            print(f"P{i}: Style='{p.style.name}', Text='{p.text[:100]}...'")
            found_any = True
    if not found_any:
        print("No automatic lists found in the entire document.")
        # Check for manual bullets
        for i, p in enumerate(doc.paragraphs):
            t = p.text.strip()
            if t.startswith('•') or t.startswith('-') or re.match(r'^\d+\.', t):
                if i < 100: # limit output
                    print(f"Manual candidate P{i}: {t[:50]}")
else:
    print("File not found.")
